from typing import TypedDict, Optional, Any
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LC_Document
from langchain_weaviate import WeaviateVectorStore as Weaviate
from langgraph.graph import StateGraph
from groq import Groq
import weaviate
from weaviate.auth import AuthApiKey
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Define State
class Mystate(TypedDict):
    text: list[str]
    split_texts: list[str]
    query: Optional[str]
    response: Optional[str]
    vectorstore: Optional[Any]

# Initialize Weaviate client and vectorstore
weaviate_client = None
vectorstore = None
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
embedding_model = HuggingFaceEmbeddings(model_name='all-MiniLM-L12-v2')

def initialize_vectorstore():
    global weaviate_client, vectorstore
    weaviate_client = weaviate.connect_to_weaviate_cloud(
        cluster_url=os.getenv("WEAVIATE_URL"),
        auth_credentials=AuthApiKey(os.getenv("WEAVIATE_API_KEY")),
        skip_init_checks=True
    )
    # Load and process document
    text = read_docx()
    split_texts = []
    for t in text:
        split_texts.extend(text_splitter.split_text(t))
    docs = [LC_Document(page_content=t) for t in split_texts]
    # Initialize vectorstore
    vectorstore = Weaviate(client=weaviate_client, index_name='LangChain', text_key='text', embedding=embedding_model)
    vectorstore.add_documents(docs)

def read_docx() -> list[str]:
    doc = Document("AI training doc.docx")
    return [para.text for para in doc.paragraphs if para.text.strip()]

def split_text(state: Mystate) -> Mystate:
    split_texts = []
    for t in state["text"]:
        split_texts.extend(text_splitter.split_text(t))
    return {
        **state,
        "split_texts": split_texts
    }

def creating_vector_store(state: Mystate) -> Mystate:
    global vectorstore
    if vectorstore is None:
        initialize_vectorstore()
    return {
        **state,
        "vectorstore": vectorstore
    }

def take_query(state: Mystate, query: str) -> Mystate:
    return {
        **state,
        "query": query
    }

def create_response(state: Mystate) -> Mystate:
    vectorstore = state["vectorstore"]
    query = state["query"]

    if not vectorstore:
        raise ValueError("Vectorstore is not initialized.")

    relevant_docs = vectorstore.similarity_search(query, k=3)
    context = "\n\n".join(doc.page_content for doc in relevant_docs)

    messages = [
        {'role': 'system', 'content': 'You are a brilliant assistant. Use the provided context to answer.'},
        {'role': 'user', 'content': f"Context:\n{context}\n\nQuestion:\n{query}"}
    ]

    client_groq = Groq(api_key=os.getenv("GROQ_API_KEY"))
    response = client_groq.chat.completions.create(
        model='llama3-8b-8192',
        messages=messages,
        temperature=0.7,
        stream=False
    )

    response_text = response.choices[0].message.content
    return {
        **state,
        "response": response_text
    }

# Build LangGraph
builder = StateGraph(Mystate)
builder.add_node("Load Document", read_docx)
builder.add_node("Split the Text", split_text)
builder.add_node("Create Vector Store", creating_vector_store)
builder.add_node("Take Query", lambda state: take_query(state, state.get("query", "")))
builder.add_node("Generate Response", create_response)

builder.set_entry_point("Load Document")
builder.add_edge("Load Document", "Split the Text")
builder.add_edge("Split the Text", "Create Vector Store")
builder.add_edge("Create Vector Store", "Take Query")
builder.add_edge("Take Query", "Generate Response")
builder.set_finish_point("Generate Response")

graph = builder.compile()

def run_pipeline(user_query: str) -> str:
    initial_state = {
        "text": [],
        "split_texts": [],
        "query": user_query,
        "response": None,
        "vectorstore": None
    }
    final_state = graph.invoke(initial_state)
    return final_state["response"]